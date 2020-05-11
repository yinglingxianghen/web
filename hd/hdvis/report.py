#coding:utf-8

import numpy as np
import math
import time
import docx
from docx import Document
from docx.shared import Inches
import time
from datetime import datetime
from datetime import timedelta
from datetime import datetime
from datetime import timedelta
from minepy import MINE
from scipy.stats import gaussian_kde
from sklearn import linear_model
import numpy as np
import random
#from hiveConnNewServer import *
from prestoConn import *
import models
import bicorr as co
import itertools as it
import predict as pdt
import healthIndex as hi

def getRawData(jznames,rawFields,startTime,endTime,dataNum):
    rawTime,postRawData = hiveRawDataTimeNew(rawFields, startTime, endTime, dataNum)
    return rawTime,postRawData
def getFengData(jznames,fengFields,startTime,endTime,dataNum):
    fengTime,postFengData = hiveDataTimeNew(fengFields, startTime, endTime, dataNum)
    return fengTime,postFengData

def hiveData(sql):
	hc = PrestoConn()
	hc.getConn()
	return hc.select(sql)
def getSql(dbName,fields,startTime,endTime,dataNum=0):
    endTimeDt = str2Dt(endTime)
    endTimeYear = str(endTimeDt.year)
    endTimeMonth = str(endTimeDt.month)
    endTimeDay = str(endTimeDt.day)
    if len(endTimeMonth) == 1:
        endTimeMonth = "0" + endTimeMonth
    if len(endTimeDay) == 1:
        endTimeDay = "0" + endTimeDay
    sql = ""
    sql_year = ""
    sql_month = ""
    sql_day = ""
    sql_num = ""
    sql_fields = fields
    if startTime == "":#如果startTime为空，则读endTime当天的数据
        sql_year = "= '" + endTimeYear + "'"
        sql_month = "= '" + endTimeMonth + "'"
        sql_day = "= '" + endTimeDay + "'"
        if dataNum > 0:
            sql_num = " limit " + str(dataNum)
        sql = "select gettime," + sql_fields + " from " + dbName + " where y " + sql_year + " and m " + sql_month + " and d " + sql_day + sql_num
        return sql
    
    startTimeDt = str2Dt(startTime)
    startTimeYear = str(startTimeDt.year)
    startTimeMonth = str(startTimeDt.month)
    startTimeDay = str(startTimeDt.day)
    sql_fields = fields
    if len(startTimeMonth) == 1:
        startTimeMonth = "0" + startTimeMonth
    if len(startTimeDay) == 1:
        startTimeDay = "0" + startTimeDay
    
    if startTimeYear == endTimeYear:
        sql_year = "= '"+ endTimeYear + "'"
    else:
        sql_year = "between '" + startTimeYear + "' and '" + endTimeYear + "'"
    if startTimeMonth == endTimeMonth:
        sql_month = "= '" + endTimeMonth + "'"
    else:
        sql_month = "between '" + startTimeMonth + "' and '" + endTimeMonth + "'"
    if startTimeDay == endTimeDay:
        sql_day = "= '" + endTimeDay + "'"
    else:
        sql_day = "between '" + startTimeDay + "' and '" + endTimeDay + "'"
    if dataNum > 0:
        sql_num = " limit "+ str(dataNum)
    sql = "select gettime," + sql_fields + " from "+ dbName + " where y "+sql_year+" and m "+sql_month+" and d "+sql_day+ " and gettime between timestamp '"\
          + startTime + "' and timestamp '" + endTime +"'" + sql_num
    return sql

def getXfsData(xfsFields,startTime,endTime,dataNum):
    dbName = "hd_middledata_all_xfs_tmp"
    slide_window = 512 * 8
    xf_len = 2049  # slide_window / 2 + 1
    __fs = 640 | 54
    __circleNo = 512
    f_step = float(__fs) / float(slide_window)
    f_data = float(__fs) / float(__circleNo)
    xf_index = float(f_data / f_step)

    x_index = []
    for i in range(xf_len):
        x_index.append(i / xf_index)

    # 获取查询字段
    sql_fields = ','.join(xfsFields)
    startTime = ''
    sql = getSql(dbName,sql_fields,startTime,endTime,dataNum)
    print sql
    sqldata = hiveData(sql)
    xfsTime = []
    dataResults = []
    for dataFrame in sqldata:
        xfsTime.append(dataFrame[0])
        dataResult = []
        for df in dataFrame[1:]:
            dr = [float(d) for d in df.split('/')]
            dataResult.append(dr)
        dataResults.append(dataResult)
    postXfsData = np.array(dataResults)

    return xfsTime, postXfsData

def hiveDataTimeNew(var, startTime, endTime, dataNum):  # 从新表获取中间数据,第一列是时间
    #dbName = 'hd_premiddledata_ffz_mean_tmp'
    dbName = 'hd_premiddledata_ffz_mean_ext_orc'
    hc = PrestoConn()
    sql_fields = ','.join(var)
    startTime = ''
    sql = getSql(dbName, sql_fields, startTime, endTime, dataNum)

    print sql
    datalist = []
    fengTime = []
    # print hc.select(sql)
    sqldatas = hc.select(sql)
    for sqldata in sqldatas:
        fengTime.append(sqldata[0])
        datalist.append(sqldata[1:])
    return fengTime,np.array(datalist)

def hiveRawDataTimeNew(BivarFields, startTime, endTime, dataNum):  # 从新表获取原始数据，,第一列是时间
    dbName = "hd_rawdata_ext_orc"
    hc = PrestoConn()
    sql_fields = ','.join(BivarFields)
    startTime = ''
    sql = getSql(dbName, sql_fields, startTime, endTime, dataNum)
    print sql
    #
    sqldatas = hc.select(sql)
    datalist = []
    rawTime = []
    for sqldata in sqldatas:
        rawTime.append(sqldata[0])
        datalist.append(sqldata[1:])
    return rawTime,np.array(datalist)

def computeCorr(x, y):
    # Calculate the point density
    xy = np.vstack([x, y])
    # z = gaussian_kde(xy)(xy)#用来画峰峰值的颜色情况
    # zz = np.floor(z * 500000)
    zz = x
    for i in zz:
        i = 50
    mine = MINE(alpha=0.6, c=15, est="mic_approx")
    mine.compute_score(x, y)
    r = np.corrcoef(x, y)[0, 1]
    
    # print_stats(mine) 最大 pearson
    return x.tolist(), y.tolist(), zz.tolist(), round(mine.mic(), 3), r

def removeDup(data):
    rows = data.shape[0]
    cols = data.shape[1]
    if rows <= 1:
        return data
    for j in range(cols):
        isDup = 1
        for i in range(rows):
            if data[i][j] != data[0][j]:
                isDup = 0
                break
        
        if isDup == 1:  # 如果重复，给这一列随机加值
            for k in range(rows):
                data[k][j] = data[k][j] + random.randint(0, 9)
    return data

def allthreholdMiddle(datamiddle):
    datamiddle = np.array(datamiddle)
    thres = []
    for i in range(datamiddle.shape[1]):
        shangx = datamiddle[:, i]
        # 均值
        
        mean = np.mean(shangx)
        
        # 无偏标准差
        std = np.std(shangx, ddof=1)
        sigma = []
        # 实际分布比例
        for i in range(3):
            xp = mean + std * (i + 1)
            xl = mean - std * (i + 1)
            k = 0
            for j in shangx:
                if j < xp and j > xl:
                    k = k + 1
            sigma.append(k / len(shangx))
        threshol = mean + 3 * std
        thres.append(round(threshol,5))
    return thres
def strToFloat(rawdata):
    floatdata = []
    for i in range(len(rawdata)):
        floatdata.append(map(eval,rawdata[i]))
    return np.array(floatdata)
def str2Dt(str):
    dt = datetime.strptime(str, '%Y-%m-%d %H:%M:%S')
    return dt
def uivarDataTable(document,jzname,categories,datalist,startTime,endTime):
    document.add_heading(u'基于大数据的轴系稳定性评估报告', 0)
    p = document.add_paragraph()
    p.add_run(u'监测起始时间:'+startTime+'\n').bold = True
    p.add_run(u'监测终止时间:'+endTime).bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    # document.add_heading(u'监测起始时间:'+startTime, level=1)
    # document.add_heading(u'监测终止时间:'+endTime, level=1)
    document.add_heading(jzname+u'机组--主要监测参数的列表', level=1)
    lenCg = len(categories)
    table = document.add_table(rows=1, cols=lenCg+1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'时刻'
    for i in range(lenCg):
        hdr_cells[i+1].text = categories[i]#i = 0是时刻
    lendatalist = 5
    for i in range(lendatalist):
        row_cells = table.add_row().cells
        for j in range(len(datalist[i])):
            row_cells[j].text = datalist[i][j]
    #document.add_page_break()
    return
def uivarTable(document,datalist,univarExceThreholds,categories,jzcompare,thre_ground,jzname):
    numcol = 7
    document.add_heading(jzname + u'机组--超标变量表', level=1)
    table = document.add_table(rows=1, cols= numcol)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u"变量名称"
    hdr_cells[1].text = u"时刻"
    hdr_cells[2].text = u"个性化阈值"
    hdr_cells[3].text = u"超过个性化阈值"
    hdr_cells[4].text = u"标准阈值"
    hdr_cells[5].text = u"超过阈值"
    hdr_cells[6].text = u"实测值"
    lendatalist = 20

    for i in range(lendatalist):
        for j in range(len(datalist[i])):
            if j == 0:
                continue;
            if univarExceThreholds[i][j-1] == 1:#已经越界，这里先用0，后面改成1表示越界, 注意datalist会多一列时间数据
                row_cells = table.add_row().cells
                row_cells[0].text = categories[j-1]
                row_cells[1].text = datalist[i][0]
                row_cells[2].text = str(jzcompare[jzname][j-1])
                row_cells[3].text = u"是"
                row_cells[4].text = str(thre_ground[j-1])
                if datalist[i][j] < thre_ground[j-1]:
                    row_cells[5].text = u"是"
                else:
                    row_cells[5].text = u"否"
                row_cells[6].text = datalist[i][j]
    return
def uivarNearTable(document,datalist,univarExceThreholds,categories,jzcompare,thre_ground,jzname):
    numcol = 5
    document.add_heading(jzname + u'机组--超标变量表', level=1)
    table = document.add_table(rows=1, cols= numcol)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u"变量名称"
    hdr_cells[1].text = u"时刻"
    hdr_cells[2].text = u"个性化阈值"
    hdr_cells[3].text = u"标准阈值"
    hdr_cells[4].text = u"实测值"
    lendatalist = 20

    for i in range(lendatalist):
        for j in range(len(datalist[i])):
            if j == 0:
                continue;
            if univarExceThreholds[i][j-1] == 0.9:#已经越界，这里先用0，后面改成1表示越界, 注意datalist会多一列时间数据
                row_cells = table.add_row().cells
                row_cells[0].text = categories[j-1]
                row_cells[1].text = datalist[i][0]
                row_cells[2].text = str(jzcompare[jzname][j-1])
                row_cells[3].text = str(thre_ground[j-1])
                row_cells[4].text = datalist[i][j]
    return
def bivarTable(document,jzname,bivarFieldpairs):
    document.add_heading(jzname + u'机组--相关关系异常变量表', level=1)
    numcol = 6
    table = document.add_table(rows=1, cols=numcol)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u"变量1"
    hdr_cells[1].text = u"变量2"
    hdr_cells[2].text = u"相关系数(标准)"
    hdr_cells[3].text = u"实测相关系数(异常)"
    hdr_cells[4].text = u"距离阈值(标准)"
    hdr_cells[5].text = u"实测距离值(异常）"

    for bivarKey in bivarFieldpairs:
        if bivarFieldpairs[bivarKey][5] < bivarFieldpairs[bivarKey][4]:#没越界
            row_cells = table.add_row().cells
            row_cells[0].text = bivarFieldpairs[bivarKey][0]
            row_cells[1].text = bivarFieldpairs[bivarKey][1]
            row_cells[2].text = str(round(bivarFieldpairs[bivarKey][2],4))
            row_cells[3].text = str(round(bivarFieldpairs[bivarKey][3],4))
            row_cells[4].text = str(round(bivarFieldpairs[bivarKey][4],4))
            row_cells[5].text = str(round(bivarFieldpairs[bivarKey][5],4))
    return
def multiTable(document,jzname,multiFieldPairs):
    #[0jzname,1fieldCom,2fieldZname,trainSamples[i].tolist(),multiTimedata[i],multiScore[i]]
    document.add_heading(jzname + u'机组--SVM变量超出表', level=1)
    lenMultiVar = 0
    for multiKey in multiFieldPairs:
        lenMultiVar = len(multiFieldPairs[multiKey][1])
        break;
    numcol = 4 + lenMultiVar
    if numcol == 4:
        return
    table = document.add_table(rows=1, cols=numcol)

    for multiKey in multiFieldPairs:
        multiFieldPair = multiFieldPairs[multiKey]
        fieldZname = multiFieldPair[2]
        row_cells = table.add_row().cells
        #说明
        row_cells[0].text = u"机组"
        row_cells[1].text = u"时间"
        for i in range(len(fieldZname)):
            row_cells[2+i].text = fieldZname[i]
        row_cells[2+len(fieldZname)].text = u"svm分类分数(>50)"
        row_cells[3+len(fieldZname)].text = u"是否超过svm边界"
        #数据
        data_cells = table.add_row().cells
        data_cells[0].text = multiFieldPair[0]
        data_cells[1].text = multiFieldPair[4]
        for i in range(len(multiFieldPair[3])):
            data_cells[2+i].text = str(round(multiFieldPair[3][i],4))
        data_cells[2+len(multiFieldPair[3])].text = str(round(multiFieldPair[5],4))
        data_cells[3+len(multiFieldPair[3])].text = u"是"
    return
def  monitTable(document,jzname,monitFields,monitNamepairs,monitRSlopes,monitSlopes):
    document.add_heading(jzname + u'机组--趋势异常变量表', level=1)
    numcol = 5
    table = document.add_table(rows=1, cols=numcol)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u"变量名"
    hdr_cells[1].text = u"时刻"
    hdr_cells[2].text = u"趋势阈值(*1000)"
    hdr_cells[3].text = u"趋势监测值(*1000)"
    hdr_cells[4].text = u"变量和趋势异常情况"

    for jzfield in monitFields:
        if jzname != monitFields[jzfield][0]:
            continue
        #lenResult = len(monitRSlopes[jzfield])
        lenResult = 20
        for j in range(lenResult):
            isExce = monitRSlopes[jzfield][j][1]
            if isExce > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = monitNamepairs[jzfield][0]#变量名
                row_cells[1].text = monitRSlopes[jzfield][j][0]#时刻
                row_cells[2].text = str(round(monitNamepairs[jzfield][3],4))#标准值
                row_cells[3].text = str(round(monitSlopes[jzfield][j][1],4))#实际值
                if isExce == 0.9:
                    row_cells[4].text = u"变量值接近阈值"
                else:
                    if isExce == 1:
                        row_cells[4].text =u"变量值超过阈值"
                    else:
                        if isExce >= 2:
                            row_cells[4].text = u"变化趋势超过阈值"
                        else:
                            row_cells[4].text = u"正常"
    return
def saveDocument(document,docPath,docName):
    document.save(docPath+docName)
    return
def hiveData(sql):
    hc = PrestoConn()
    hc.getConn()
    return hc.select(sql)
def getSql(dbName, fields, startTime, endTime, dataNum=0):
    endTimeDt = str2Dt(endTime)
    endTimeYear = str(endTimeDt.year)
    endTimeMonth = str(endTimeDt.month)
    endTimeDay = str(endTimeDt.day)
    if len(endTimeMonth) == 1:
        endTimeMonth = "0" + endTimeMonth
    if len(endTimeDay) == 1:
        endTimeDay = "0" + endTimeDay
    sql = ""
    sql_year = ""
    sql_month = ""
    sql_day = ""
    sql_num = ""
    sql_fields = fields
    if startTime == "":  # 如果startTime为空，则读endTime当天的数据
        sql_year = "= '" + endTimeYear + "'"
        sql_month = "= '" + endTimeMonth + "'"
        sql_day = "= '" + endTimeDay + "'"
        if dataNum > 0:
            sql_num = " limit " + str(dataNum)
        sql = "select gettime," + sql_fields + " from " + dbName + " where y " + sql_year + " and m " + sql_month + " and d " + sql_day + sql_num
        return sql
    
    startTimeDt = str2Dt(startTime)
    startTimeYear = str(startTimeDt.year)
    startTimeMonth = str(startTimeDt.month)
    startTimeDay = str(startTimeDt.day)
    sql_fields = fields
    if len(startTimeMonth) == 1:
        startTimeMonth = "0" + startTimeMonth
    if len(startTimeDay) == 1:
        startTimeDay = "0" + startTimeDay
    
    if startTimeYear == endTimeYear:
        sql_year = "= '" + endTimeYear + "'"
    else:
        sql_year = "between '" + startTimeYear + "' and '" + endTimeYear + "'"
    if startTimeMonth == endTimeMonth:
        sql_month = "= '" + endTimeMonth + "'"
    else:
        sql_month = "between '" + startTimeMonth + "' and '" + endTimeMonth + "'"
    if startTimeDay == endTimeDay:
        sql_day = "= '" + endTimeDay + "'"
    else:
        sql_day = "between '" + startTimeDay + "' and '" + endTimeDay + "'"
    if dataNum > 0:
        sql_num = " limit " + str(dataNum)
    sql = "select gettime," + sql_fields + " from " + dbName + " where y " + sql_year + " and m " + sql_month + " and d " + sql_day + " and gettime >= timestamp '" \
          + startTime + "' and gettime <= timestamp '" + endTime + "'" + sql_num
    return sql
def getRawFields(fields):
    rawFields = []
    for field in fields:
        pos = field.rfind('_')
        if pos == -1:
            rawFields.append(field)
        else:
            rawFields.append(field[:pos])
    return rawFields
def getZname(univarFields):
    fieldZname = {}
    categories = []  # 变量种类,中文表示
    thre_ground = []  # 标准阈值，对于每个变量早定好了
    
    for field in univarFields:
        if field != '':
            fieldZname[field] = models.Middlevariable.objects.filter(fieldname=field)[0].comments
            categories.append(models.Middlevariable.objects.filter(fieldname=field)[0].comments)
            thre_ground.append(models.Middlevariable.objects.filter(fieldname=field)[0].hech)
    return fieldZname,categories,thre_ground

def univarModule(jznames,univarFields,startTime, endTime, data_num,
                 userid,activepower,rotatespeed1,rotatespeed2,rotatespeed3):
    jzcompare = {}
    fengTime,datalist = hiveDataTimeNew(univarFields, startTime, endTime, data_num)  # 根据终点时刻选择数据,第一列表示该数据的时刻,所有跟数据有关要注意
    dataTime = fengTime
    datamiddle = datalist.tolist()
    datalist = []
    for i in range(len(dataTime)):
        d = []
        d.append(dataTime[i])
        for j in range(len(datamiddle[i])):
            d.append(datamiddle[i][j])
        datalist.append(d)
    
    # 通过查询获取个性化阈值,保存在jzcompare
    print("bivar...\n")
    for jzname in jznames:  # 根据机组，变量从models中获取个性化阈值
        thre_personalize = []
        for field in univarFields:
            md = models.Personalizedthreshold.objects.filter(userid=userid, aliasname=field, jzname=jzname,
                                                             activepower=activepower,
                                                             rotatespeed1=rotatespeed1, rotatespeed2=rotatespeed2,
                                                             rotatespeed3=rotatespeed3)
            if (md.count() != 0):
                thre_personalize.append(md[0].threshold)
        jzcompare[jzname] = thre_personalize
    univarExceThreholds = []
    for jzname in jznames:  # 根据机组，变量从models中获取个性化阈值
        for i in range(len(datalist)):
            isExceed = 0
            exceThrehold = []
            for j in range(len(datalist[i])):
                if j == 0:
                    continue
                if (len(jzcompare[jzname]) == 0):
                    continue
                datalist[i][j] = str(float(datalist[i][j]) * 1.0)  # 将数据乘以1.5
                if float(datalist[i][j]) < float(jzcompare[jzname][j - 1]) * 0.9:  # 变量没越界，datalist比jzcompare多一列
                    exceThrehold.append(0)
                else:
                    if float(datalist[i][j]) > float(jzcompare[jzname][j - 1]):  # 变量近似越界
                        exceThrehold.append(1)
                    else:
                        exceThrehold.append(0.9)
            jzfieldid = jzname + univarFields[j - 1]
            jzid = jzname + " " + datalist[i][0]
            univarExceThreholds.append(exceThrehold)  # jz号机组越界情况
    
    # 相关性分析
    return jzcompare,datalist,univarExceThreholds,datalist,dataTime,datamiddle
def getJzcompare(jznames,univarFields,userid,activepower,rotatespeed1,rotatespeed2,rotatespeed3):
    jzcompare = {}  # 个性化阈值
    
    print("bivar...\n")
    for jzname in jznames:  # 根据机组，变量从models中获取个性化阈值
        thre_personalize = []
        for field in univarFields:
            md = models.Personalizedthreshold.objects.filter(userid=userid, aliasname=field, jzname=jzname,
                                                             activepower=activepower,
                                                             rotatespeed1=rotatespeed1, rotatespeed2=rotatespeed2,
                                                             rotatespeed3=rotatespeed3)
            if (md.count() != 0):
                thre_personalize.append(md[0].threshold)
        jzcompare[jzname] = thre_personalize
    return jzcompare
def bivarModule(jznames,bivarFields,bivarRawFields,startTime,endTime,data_num,userid,
                mainField,fieldZname):
    
    if mainField == '':
        mainCom = [x for x in it.combinations(bivarFields, 2)]
    else:
        if mainField in bivarFields:
            bivarFields.remove(mainField)
        mainCom = [[mainField, f] for f in bivarFields]  # 进行组合
        bivarFields.insert(0, mainField)  # mainfield在第一个位置
    bivarFieldZname = {}
    for bf in bivarFields:
        bivarFieldZname[bf] = models.Middlevariable.objects.filter(fieldname=bf)[0].comments
    bivarFieldsIdx = {}  # 变量对应在rawdata中的序号
    for i in range(len(bivarFields)):
        bivarFieldsIdx[bivarFields[i]] = i
    
    # 读取数据
    rawTime,rawData = hiveRawDataTimeNew(bivarRawFields, startTime, endTime, data_num)  # 第一列为时间
    fentTime,fengData = hiveDataTimeNew(bivarFields, startTime, endTime, data_num)  # 第一列为时间

    
    # 从数据库获取时间以便于读取数据
    md = models.Bivarvariable.objects.filter(userid=userid)
    dbStartTime = ""
    dbEndTime = ""
    
    if md.count != 0:
        dbStartTime = md[0].starttime
        dbEndTime = md[0].endtime
        dbDataNum = md[0].datanum
    else:
        dbStarTime = startTime
        dbEndTime = endTime
        dbDataNum = 200
    dbrawTime,dbrawdata = hiveRawDataTimeNew(bivarRawFields, dbStartTime, dbEndTime, dbDataNum)  # 第一列为时间
    
    fengDbTime,fengDbData = hiveDataTimeNew(bivarFields, dbStartTime, dbEndTime, dbDataNum)  # 第一列为时间
    
    # 初始化相关数据
    datapair = {}
    dbdatapair = {}
    fengdatapair = {}  # 峰峰值数据
    fengdbdatapair = {}
    bivarFengFieldPairs = {}  # 峰峰值相关数据
    bivarFieldPairs = {}  # 原始相关数据
    fields_no = rawData.shape[1]
    rotatespeed = 0  # 转速
    load = 0  # 负荷
    activepower = 0  # 有功功率
    
    print("bivar......\n")
    for jzname in jznames:
        for i in range(0, len(mainCom)):  # 第一列是主变量
            fieldpair = ''.join(mainCom[i])
            fieldpair = jzname + fieldpair
            aliasname1 = mainCom[i][0]
            aliasname2 = mainCom[i][1]
            idx1 = bivarFieldsIdx[aliasname1]
            idx2 = bivarFieldsIdx[aliasname2]
            dbbivarr = models.Bivarvariable.objects.filter(aliasname1=aliasname1, aliasname2=aliasname2)
            if dbbivarr.count() == 0:  #
                continue
            bivarr = dbbivarr[0].bivarr
            if (abs(bivarr) < 0):  # 选的这两个维度健康数据相关系数没有大于0.8,认为是没有线性相关性的，直接跳过
                continue
            # 对第0列和第i列进行相关性判断，算这两列x,y的点到直线的距离
            x, y, z, mic, r = co.computeCorrScale(rawData[:, idx1], rawData[:, idx2], 1.0)
            dbx, dby, dbz, dbmic, dbr = co.computeCorr(dbrawdata[:, idx1], dbrawdata[:, idx2])
            fengx, fengy, fengz, fengmic, fengr = co.computeCorr(fengData[:, idx1], fengData[:, idx2])
            fengdbx, fengdby, fengdbz, fengdbmic, fengdbr = co.computeCorrScale(fengDbData[:, idx1],
                                                                                fengDbData[:, idx2], 1.0)
            dbcoef = dbbivarr[0].coef
            dbintercept = dbbivarr[0].intercept
            dbdistthre = dbbivarr[0].distthre
            dbbivarr = dbbivarr[0].bivarr  # 线性系数阈值
            # dbdistthre = dbbivarr[0].bivarr
            k = dbcoef
            b = dbintercept
            dists = []
            for j in range(len(x)):
                dist = abs(k * x[j] - y[j] + b) / ((k * k + 1) ** 0.5)
                dists.append(dist)
            distmean = np.mean(dists)  # 距离均值
            distret = 0  # 没越界
            if distmean > dbdistthre:  # 距离已经越界
                distret = 0
            coef, intercept, score = co.linearRegression(x, y)  # 获取两列数据拟合值
            
            dbdatapair[fieldpair] = [dbx, dby, dbz]
            datapair[fieldpair] = [x, y, z]
            fengdatapair[fieldpair] = [fengx, fengy, fengz]
            fengdbdatapair[fieldpair] = [fengdbx, fengdby, fengdbz]
            # 判断是否符合线性相关
            namepair = [fieldZname[aliasname1],
                        fieldZname[aliasname2], mic, r, dbdistthre, distmean, jzname, dbmic, dbr]
            fengnamepair = [fieldZname[aliasname1],
                            fieldZname[aliasname2], fengmic, fengr, jzname, fengdbmic, fengdbr]
            bivarFengFieldPairs[fieldpair] = fengnamepair
            bivarFieldPairs[fieldpair] = namepair
    return bivarFieldPairs,bivarFengFieldPairs,dbdatapair,datapair,fengdatapair,fengdbdatapair
def monitModule(jznames,fieldZname,monitSingleFields,datamiddle,dataTime,univarExceThreholds):
    # 处理变量
    monitMiddleData = datamiddle
    monitMiddleData = np.array(monitMiddleData)
    monitDatetime = pdt.str2Dt(dataTime)
    
    # 在report中先要进行阈值判断
    rotatespeed = 0
    load = 0
    activepower = 0
    slopejz = {}
    timeInterval = 1  # 时间间隔 1min
    namepairs = {}
    monitNamepairs = {}  # 保留那些故障的维度
    monitDatas = {}  # 原始数据，时间
    monitSlopes = {}  # 斜率，时间
    monitRSlopes = {}  # 越界结果，1为变量阈值越界，2为斜率阈值越界
    monitFields = {}  # 机组，变量，中文名，机组变量
    monitIntercepts = {}
    monitSegmentDatas = {}  # 以jzname+filed+time作为id,保存斜率变量超标的数据
    univarExceThreholds = np.array(univarExceThreholds)
    
    # 斜率阈值都乘以1000
    print("monit......\n")
    factor = 1000
    for jzname in jznames:
        isExce = 0
        for i in range(len(monitSingleFields)):
            # datadb, X_slopedb, slopedb = pdt.timeIntervalSlope(datetimedata, rawdatas[:, i], timeInterval)#时间段为timeInterval的斜率
            # 计算出这一列的斜率值
            field = monitSingleFields[i]
            zname = fieldZname[field]
            md = models.Monitorvariable.objects.filter(jzname=jzname, rotatespeed=rotatespeed, load=load,
                                                       activepower=activepower, aliasname=monitSingleFields[i],
                                                       timeinterval=timeInterval)
            if md.count() == 0:  # 数据库中还未保存
                continue
            
            thre_slope = md[0].mid_val  # mid_val = mean + 2 * std，斜率乘以1000,做判断
            alert_slope = md[0].alert_val  # 用来画图的
            monitXSlope, monitXIntercept = pdt.timeSlope(monitDatetime,
                                                         univarExceThreholds[:, i],
                                                         monitMiddleData[:, i],
                                                         timeInterval, thre_slope, jzname, zname, field,
                                                         monitSegmentDatas, alert_slope, factor)  # 一列数值
            monitRSlope, monitExce, monitData, monitSlope, monitIntercept = pdt.judgeSlope(monitDatetime,
                                                                                           univarExceThreholds[:,
                                                                                           i],
                                                                                           monitMiddleData[:, i],
                                                                                           monitXSlope,
                                                                                           monitXIntercept,
                                                                                           thre_slope)
            min_val, mid_val, max_val, alert_val, mean, std = pdt.meanstd(monitXSlope)
            jzfield = jzname + monitSingleFields[i]  # 机组加变量名
            field = monitSingleFields[i]
            zname = fieldZname[monitSingleFields[i]]  # 中文名
            if monitExce == 1:  # 出现故障
                monitFields[jzfield] = [jzname, field, zname, jzfield]
                monitNamepairs[jzfield] = [zname, thre_slope, min_val, mid_val, max_val, alert_val]  # 仅保留名称即可,
                monitDatas[jzfield] = monitData
                monitSlopes[jzfield] = monitSlope
                monitIntercepts[jzfield] = monitIntercept
                monitRSlopes[jzfield] = monitRSlope
                
    return monitFields,monitNamepairs,monitDatas,monitSlopes,monitIntercepts,monitRSlopes,monitSegmentDatas

def resultBivar(jznames,fields,startTime,endTime,data_num,userid,fieldZname):
    # 变量维度
    BivarFields = fields  # 对变量按照英文字母进行排序
    multiFieldsIdx = {}  # 变量对应在rawdata中的序号
    for i in range(len(BivarFields)):
        multiFieldsIdx[BivarFields[i]] = i
    
    mainCom = [x for x in it.combinations(BivarFields, 2)]  # 进行22组合
    lenCom = len(mainCom)
    
    # modelresult处理峰峰值数据# 处理原始数据,保留原始数据中相关系数大于0.8的情况
    print("bivar...")
    bivarFengDataNum = data_num
    bivarRawDataNum = data_num
    
    # 原始值
    BivarRawFields = getRawFields(fields)
    bivarRawTime,bivarRawData = hiveRawDataTimeNew(BivarRawFields, startTime, endTime,
                                          bivarRawDataNum)  # startTime为空，取当天前datanum条数据，

    
    # 峰峰值中间数据读取
    datatime,bivarFengData = hiveDataTimeNew(BivarFields, startTime, endTime, bivarFengDataNum)  # 读取前2000个，速度比较快
    
    datamiddle = bivarFengData
    
    # 声明显示相关变量
    bivarRawDataPairs = {}
    bivarRawFieldPairs = {}
    bivarFengDataPairs = {}
    bivarFengFieldPairs = {}
    fields_no = bivarFengData.shape[1]
    numNoBicorr = 0
    rotatespeed = 0
    load = 0
    activepower = 0
    # modelresult从数据库中读取数据进行显示
    for jzname in jznames:
        for i in range(0, lenCom):
            # 原始数据相关性
            fieldpair = ''.join(mainCom[i])
            fieldpair = jzname + fieldpair
            aliasname1 = mainCom[i][0]
            aliasname2 = mainCom[i][1]
            md = models.Bivarvariable.objects.filter(userid=userid, jzname=jzname, aliasname1=aliasname1,
                                                     aliasname2=aliasname2)
            if md.count() == 0:
                continue
            idx1 = multiFieldsIdx[mainCom[i][0]]
            idx2 = multiFieldsIdx[mainCom[i][1]]
            # raw
            rawx, rawy, rawz, rawmic, rawr = co.computeCorr(bivarRawData[:, idx1], bivarRawData[:, idx2])
            distthre = md[0].distthre
            distmean = 0  # 这里貌似没用
            bivarRawDataPairs[fieldpair] = [rawx, rawy, rawz]
            bivarRawFieldPairs[fieldpair] = [fieldZname[aliasname1],
                                             fieldZname[aliasname2], rawmic, rawr, distthre, distmean, jzname]
            # feng
            fengx, fengy, fengz, fengmic, fengr = co.computeCorr(bivarFengData[:, idx1], bivarFengData[:, idx2])
            fengmic = md[0].fengmic
            fengr = md[0].fengr
            bivarFengDataPairs[fieldpair] = [fengx, fengy, fengz]
            bivarFengFieldPairs[fieldpair] = [fieldZname[aliasname1],
                                              fieldZname[aliasname2], fengmic, fengr, jzname]
    
    return bivarRawDataPairs,bivarRawFieldPairs,bivarFengDataPairs,bivarFengFieldPairs,datamiddle,datatime

def resultIsof(startTime,endTime,isofNum =1000,isofName='healthModelxxx'):
    isofTime, isofZ = hi.health_index(startTime, endTime, isofNum, isofName)
    return isofTime, isofZ
def resultMonit(jznames,fields,datamiddle,datatime,timeIntervals,fieldZname):
    # modelresult处理变量
    monitSingleFields = fields
    # 读取数据
    monitMiddleData = datamiddle
    monitMiddleData = np.array(monitMiddleData)
    monitDatetime = pdt.str2Dt(datatime)
    
    # 在report中先要进行阈值判断
    rotatespeed = 0
    load = 0
    activepower = 0
    slopejz = {}
    # timeInterval = 1  # 时间间隔 1min
    namepairs = {}
    monitNamepairs = {}  # 保留那些故障的维度
    monitDatas = {}  # 原始数据，时间
    monitSlopes = {}  # 斜率，时间
    monitRSlopes = {}  # 越界结果，1为变量阈值越界，2为斜率阈值越界
    monitFields = {}  # 机组，变量，中文名，机组变量
    monitIntercepts = {}
    monitSegmentDatas = {}  # 以jzname+filed+time作为id,保存斜率变量超标的数据
    
    univarExceThreholds = []
    for i in range(len(monitMiddleData)):
        exceThrehold = []
        for j in range(len(monitMiddleData[i])):
            exceThrehold.append(0)
        univarExceThreholds.append(exceThrehold)  # jz号机组越界情况
    univarExceThreholds = np.array(univarExceThreholds)
    # modelresult斜率阈值都乘以1000
    print("monit......\n")
    factor = 1000
    for jzname in jznames:
        isExce = 0
        for i in range(len(monitSingleFields)):
            for timeInterval in timeIntervals:
                
                # datadb, X_slopedb, slopedb = pdt.timeIntervalSlope(datetimedata, rawdatas[:, i], timeInterval)#时间段为timeInterval的斜率
                # 计算出这一列的斜率值
                
                field = monitSingleFields[i]
                zname = fieldZname[field]
                md = models.Monitorvariable.objects.filter(jzname=jzname, rotatespeed=rotatespeed, load=load,
                                                           activepower=activepower, aliasname=monitSingleFields[i],
                                                           timeinterval=timeInterval)
                if md.count() == 0:  # 数据库中还未保存
                    continue
                min_val = md[0].min_val
                mid_val = md[0].mid_val
                max_val = md[0].max_val
                alert_val = md[0].alert_val
                mean = md[0].mean  # mean = min_val
                std = md[0].std
                thre_slope = md[0].mid_val
                jzfield = timeInterval + '_' + jzname + '_' + monitSingleFields[i]  # 机组加变量名
                monitNamepairs[jzfield] = [zname, thre_slope, min_val, mid_val, max_val, alert_val, mean,
                                           std]  # 仅保留名称即可,
                monitFields[jzfield] = [jzname, field, zname, jzfield, timeInterval]
    return monitNamepairs,monitFields